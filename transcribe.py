import speech_recognition as sr
from os import path, makedirs
from pydub import AudioSegment
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Define directory paths
input_dir = "raw_audio"
conversion_dir = "conversions"
transcription_dir = "python_transcriptions"

# Prompt for the file name
file_name = input("Enter the name of the audio file (including extension): ")

# Define input file path
input_file_path = path.join(input_dir, file_name)

# Check if the input file exists
if not path.isfile(input_file_path):
    print(f"File '{file_name}' does not exist in the '{input_dir}' directory.")
    exit()

# Create the conversion directory if it doesn't exist
makedirs(conversion_dir, exist_ok=True)

# Convert the audio file to WAV format
sound = AudioSegment.from_file(input_file_path)
output_file_name = file_name.rsplit(".", 1)[0] + ".wav"
output_file_path = path.join(conversion_dir, output_file_name)
sound.export(output_file_path, format="wav")
print(f"Audio file converted to WAV: {output_file_path}")

# Create the transcription directory if it doesn't exist
makedirs(transcription_dir, exist_ok=True)

# Define the transcription file path
transcription_file_name = file_name.rsplit(".", 1)[0] + ".docx"
transcription_file_path = path.join(transcription_dir, transcription_file_name)

# Transcribe the audio file
AUDIO_FILE = output_file_path

# Use the audio file as the audio source
r = sr.Recognizer()
with sr.AudioFile(AUDIO_FILE) as source:
    audio = r.record(source)  # Read the entire audio file

    # Perform speech recognition
    try:
        transcript = r.recognize_google(audio)

        # Create a new Word document
        doc = Document()
        doc.add_paragraph(transcript)

        # Apply formatting to the transcript
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(12)
                font.bold = False

        # Save the transcript to a Word document
        doc.save(transcription_file_path)
        print(f"Transcript exported to: {transcription_file_path}")
    except sr.UnknownValueError:
        print("Speech recognition could not understand audio")
    except sr.RequestError as e:
        print("Error during speech recognition: {0}".format(e))
